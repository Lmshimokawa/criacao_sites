"""
Extrai conteúdo de .docx e gera Markdown (parágrafos, títulos, listas).
Uso: python docx_to_md.py <arquivo.docx>
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def style_to_heading_level(style_name):
    if not style_name:
        return 0
    s = (style_name or "").lower()
    if "heading 1" in s or "título 1" in s:
        return 1
    if "heading 2" in s or "título 2" in s:
        return 2
    if "heading 3" in s or "título 3" in s:
        return 3
    if "heading 4" in s or "título 4" in s:
        return 4
    return 0


def get_paragraph_style(doc, paragraph):
    """Nome do estilo do parágrafo."""
    if paragraph.style and paragraph.style.name:
        return paragraph.style.name
    return None


def paragraph_to_md(doc, paragraph):
    """Converte um parágrafo docx em linha(s) markdown."""
    style_name = get_paragraph_style(doc, paragraph)
    level = style_to_heading_level(style_name)
    text = paragraph.text.strip()
    if not text:
        return ""

    # Título
    if level >= 1:
        return "\n" + "#" * level + " " + text + "\n"

    # Lista (bullet/numbering)
    if paragraph._element.xpath(".//w:numPr"):
        return "- " + text + "\n"
    # Indentação forte pode indicar sublista
    if paragraph.paragraph_format.left_indent and paragraph.paragraph_format.left_indent.pt and paragraph.paragraph_format.left_indent.pt > 0:
        return "  - " + text + "\n"

    return text + "\n\n"


def docx_to_md(docx_path: Path) -> str:
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        line = paragraph_to_md(doc, para)
        if line:
            lines.append(line)
    # Tabelas: simples por linha
    for table in doc.tables:
        lines.append("\n")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |\n")
        lines.append("\n")
    return "".join(lines).strip()


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python docx_to_md.py <arquivo.docx> [saida.md]", file=sys.stderr)
        sys.exit(1)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"Arquivo não encontrado: {path}", file=sys.stderr)
        sys.exit(1)
    out = docx_to_md(path)
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else path.with_suffix(".md")
    out_path.write_text(out, encoding="utf-8")
    print(f"Escrito: {out_path}", file=sys.stderr)
